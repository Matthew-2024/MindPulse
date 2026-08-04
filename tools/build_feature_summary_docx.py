from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "心晴MindPulse_现有功能总结.docx"
IMG_HOME = ROOT / "docs" / "review-evidence" / "02-home-status.png"
IMG_RISK = ROOT / "docs" / "review-evidence" / "05-high-risk-feedback.png"

# compact_reference_guide preset + named MindPulse brand overrides.
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "E8EEF5"
TEAL = "1E6B5C"
TEAL_DARK = "17433C"
TEAL_LIGHT = "EAF6F2"
MINT = "BFE3D8"
INK = "23312F"
MUTED = "66736F"
RISK = "9B1C1C"
CAUTION = "7A5A00"
WHITE = "FFFFFF"


def set_run_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None,
                 color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D6DEDB", size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, display, end])
    set_run_font(run, size=9, color=MUTED)


def add_numbering(document, kind="bullet"):
    numbering = document.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abs, default=-1) + 1
    num_id = max(existing_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "271")
    ppr.append(ind)
    lvl.append(ppr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def keep_together(paragraph, with_next=False):
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.widow_control = True
    if with_next:
        paragraph.paragraph_format.keep_with_next = True


def add_bullet(doc, bullet_num_id, text, bold_prefix=None):
    p = doc.add_paragraph()
    apply_num(p, bullet_num_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, color=INK, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)
    keep_together(p)
    return p


def add_numbered(doc, num_id, title, text):
    p = doc.add_paragraph()
    apply_num(p, num_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r1 = p.add_run(title + "：")
    set_run_font(r1, size=11, color=TEAL_DARK, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=11, color=INK)
    keep_together(p)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    keep_together(p, with_next=True)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=11, color=INK, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)
    keep_together(p)
    return p


def add_callout(doc, label, text, color=TEAL_DARK, fill=TEAL_LIGHT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.line_spacing = 1.2
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), color)
    borders.append(left)
    ppr.append(borders)
    r1 = p.add_run(label + "  ")
    set_run_font(r1, size=10.5, color=color, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    keep_together(p)
    return p


def fill_table_text(table, rows, header=True, widths=(2700, 6660)):
    for row_idx, row_data in enumerate(rows):
        cells = table.rows[row_idx].cells
        for col_idx, value in enumerate(row_data):
            p = cells[col_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(value)
            is_header = header and row_idx == 0
            set_run_font(run, size=10 if is_header else 9.5,
                         color=WHITE if is_header else INK, bold=is_header)
            if is_header:
                set_cell_shading(cells[col_idx], TEAL)
    if header:
        set_repeat_table_header(table.rows[0])
    set_table_geometry(table, list(widths))
    set_table_borders(table)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def configure_section(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("心晴 MindPulse  |  现有功能总结")
    set_run_font(r, size=9, color=MUTED, bold=True)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    r = fp.add_run("截至 2026-07-27  |  第 ")
    set_run_font(r, size=9, color=MUTED)
    add_page_field(fp)
    r = fp.add_run(" 页")
    set_run_font(r, size=9, color=MUTED)


def add_title_block(doc):
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(5)
    r = kicker.add_run("PRODUCT CAPABILITY BRIEF")
    set_run_font(r, size=9.5, color=TEAL, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run("心晴 MindPulse 现有功能总结")
    set_run_font(r, size=27, color=TEAL_DARK, bold=True)
    keep_together(title, with_next=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run("大学生心理压力早期觉察与端侧轻干预系统原型")
    set_run_font(r, size=13, color=MUTED)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(14)
    r = meta.add_run("版本口径：当前仓库实现  |  整理日期：2026 年 7 月 27 日  |  主交付：Web 原型 + SwiftUI 原型")
    set_run_font(r, size=9.5, color=MUTED)


def build():
    doc = Document()
    configure_styles(doc)
    configure_section(doc)
    bullet_id = add_numbering(doc, "bullet")
    number_id = add_numbering(doc, "decimal")
    add_title_block(doc)

    add_callout(
        doc,
        "一句话概括",
        "心晴把“匿名记录—状态解释—轻量行动—效果回写—必要时求助”串成一个端侧闭环；它用于早期觉察和求助转接，不进行心理诊断，也不替代专业咨询。",
    )

    add_heading(doc, "1. 产品定位与当前形态", 1)
    add_body(doc, "目标用户：有压力、焦虑、低落或睡眠紊乱，但尚未主动求助的大学生。产品重点不是给出医学结论，而是帮助用户看见自己的状态、完成一个低负担行动，并在需要时更容易开口求助。", "目标用户：")

    snapshot_rows = [
        ("维度", "当前实现"),
        ("主要载体", "Web 原型包，通过入口 HTML 运行，需保留同目录 src 脚本与资源；另有独立 SwiftUI 原型代码。"),
        ("数据方式", "Web 使用 localStorage；iOS 原型使用 UserDefaults。均以本机匿名档案隔离记录。"),
        ("判断方式", "MindPulse-RISE 可解释规则 + 个人基线 + Safety Gate 风险闸门。"),
        ("服务边界", "早期觉察、轻干预、求助表达与资源转接；不诊断、不治疗、不自动上报。"),
    ]
    table = doc.add_table(rows=len(snapshot_rows), cols=2)
    fill_table_text(table, snapshot_rows, header=True, widths=(2400, 6960))

    add_heading(doc, "2. 用户使用闭环", 1)
    steps = [
        ("进入匿名档案", "创建或选择本地匿名档案，无需手机号、学号或实名信息。"),
        ("记录当下状态", "填写情绪、睡眠、步数、精力、社交连接度和一句备注；同一天可保存多条带时间戳的即时记录。"),
        ("获得可解释判断", "系统计算 RISE 恢复指数，展示分数拆解、个人基线偏移、风险证据和推荐依据。"),
        ("完成一个微行动", "按状态进入呼吸、散步、专注、记录、睡前放松或联系朋友等轻干预。"),
        ("回写与个性化", "完成情况及前后分数变化写回当前档案，下一次优先参考对该用户更有效的动作。"),
        ("复盘或求助", "在趋势页查看日报/周报；中高风险时强化联系他人，高风险时 Safety Gate 停止普通自助主路径。"),
    ]
    for title, text in steps:
        add_numbered(doc, number_id, title, text)

    add_heading(doc, "3. 已实现功能模块", 1)

    add_heading(doc, "3.1 匿名档案与本地身份隔离", 2)
    add_bullet(doc, bullet_id, "支持创建、进入、切换和删除多个匿名档案；每个档案都有本地匿名 ID。")
    add_bullet(doc, bullet_id, "不同档案的记录、个人基线、干预完成情况和推荐反馈相互隔离。")
    add_bullet(doc, bullet_id, "不要求手机号、学号或实名信息；正式校园试点的身份认证仍未接入。")

    add_heading(doc, "3.2 即时记录与多信号输入", 2)
    add_bullet(doc, bullet_id, "可记录情绪类型、睡眠小时、步数、精力、社交连接度及文本备注。")
    add_bullet(doc, bullet_id, "同一天可连续追加多条即时记录，保留日内情绪和节奏变化，而不是只覆盖当天最后一次状态。")
    add_bullet(doc, bullet_id, "Web 端支持手动输入真实睡眠、步数和连接度；当前展示数据仍以手动/演示输入为主。")

    add_heading(doc, "3.3 RISE 恢复指数与个人基线", 2)
    add_bullet(doc, bullet_id, "RISE 汇总 Rhythm 节奏、Interaction 连接、Self-report 自评和 Engagement 干预反馈，页面对应情绪、睡眠、活动、连接与干预完成分。")
    add_bullet(doc, bullet_id, "首页展示当前恢复指数，并可进入判断详情查看每项分数、风险依据、推荐路径和数据来源。")
    add_bullet(doc, bullet_id, "根据当前档案近几天历史形成个人基线，判断睡眠、活动、连接和情绪是否偏离自己的正常水平，不与他人比较。")

    add_heading(doc, "3.4 风险分级与 Safety Gate", 2)
    add_bullet(doc, bullet_id, "提供稳定观察、普通波动、中度关注和高风险四级风险结果。")
    add_bullet(doc, bullet_id, "规则覆盖危机关键词、连续低睡眠、连续负面状态、持续社交断联及部分否定短语，降低明显误触发。")
    add_bullet(doc, bullet_id, "高风险时停止普通自助建议，优先展示求助入口、热线和可信任成人；风险等级仅用于决定路径，不用于诊断。")

    add_heading(doc, "3.5 自适应陪伴与轻干预工具", 2)
    add_bullet(doc, bullet_id, "根据情绪、风险、个人基线和历史反馈生成三步左右的陪伴路线，并逐步推进。")
    add_bullet(doc, bullet_id, "已实现 4-4-4-4 方形呼吸、10 分钟专注、计时散步、自由记录、睡前放松和联系朋友等行动。")
    add_bullet(doc, bullet_id, "睡前放松内置由 Web Audio 生成的呼吸底噪、雨声和森林氛围，可切换声音、调节音量、静音和计时。")
    add_bullet(doc, bullet_id, "记录干预完成次数及前后恢复指数变化，形成档案内的推荐反馈学习。")

    add_heading(doc, "3.6 趋势、日报与周报", 2)
    add_bullet(doc, bullet_id, "趋势页汇总平均睡眠、低睡眠天数、总步数、情绪起伏轨迹和明显转折点。")
    add_bullet(doc, bullet_id, "今日多条记录自动形成时间线和日报；近 7 天记录自动生成周报。")
    add_bullet(doc, bullet_id, "日报/周报包含平均恢复指数、低睡眠日、Safety Gate 触发情况等，并随 JSON 一并导出。")

    add_heading(doc, "3.7 轻自查问卷", 2)
    add_bullet(doc, bullet_id, "Web 首页可进入 PHQ-9、GAD-7 和压力感自查，支持逐题完成、计分、等级解释和历史结果展示。")
    add_bullet(doc, bullet_id, "问卷结果由规则逻辑解释，属于自我觉察辅助，不作为临床诊断。")

    add_heading(doc, "3.8 待办清单与低压力行动拆分", 2)
    add_bullet(doc, bullet_id, "支持新增、展开、完成、重新打开和删除待办事项。")
    add_bullet(doc, bullet_id, "待办页强调“先做最小一步”，打开事项时显示舒缓提示，并可直接转到陪伴或压力自查。")

    add_heading(doc, "3.9 求助表达与资源入口", 2)
    add_bullet(doc, bullet_id, "根据当前风险推荐优先联系对象，如朋友、家人、老师或辅导员。")
    add_bullet(doc, bullet_id, "用户可选择联系对象、希望对方如何回应、表达风格，并补充情境；系统生成可复制的求助话术草稿。")
    add_bullet(doc, bullet_id, "求助页展示校园/热线资源和拨号入口；系统不会替用户发送消息，也不会自动上报。")

    add_heading(doc, "3.10 设置、隐私与数据控制", 2)
    add_bullet(doc, bullet_id, "设置页集中管理匿名档案、个体化模型说明、隐私边界、记录来源和验证入口。")
    add_bullet(doc, bullet_id, "支持导出 JSON 与删除当前本地数据；导出内容包含匿名档案、原始记录、风险结果、干预完成情况、日报和周报。")
    add_bullet(doc, bullet_id, "提供每日轻提醒、睡前放松提醒的原型开关；当前未接入系统级后台通知。")
    add_bullet(doc, bullet_id, "Web 已配置 manifest、图标和 Service Worker，具备基础 PWA 结构。")

    # Named pagination override: keep the visual evidence section together.
    doc.add_page_break()
    add_heading(doc, "4. 关键界面示例", 1)
    add_body(doc, "以下截图来自仓库现有的自动化审查证据：左图展示日常状态与下一步，右图展示高风险记录触发 Safety Gate 后的拦截反馈。")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    if IMG_HOME.exists():
        home_pic = p.add_run().add_picture(str(IMG_HOME), width=Inches(2.18))
        home_pic._inline.docPr.set("descr", "心晴首页：恢复指数、风险等级、Safety Gate 状态与开始第一步按钮")
        home_pic._inline.docPr.set("title", "日常判断首页")
    p.add_run("      ")
    if IMG_RISK.exists():
        risk_pic = p.add_run().add_picture(str(IMG_RISK), width=Inches(2.18))
        risk_pic._inline.docPr.set("descr", "高风险记录反馈：暂停普通练习并优先打开求助入口")
        risk_pic._inline.docPr.set("title", "Safety Gate 高风险反馈")
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(6)
    r = caption.add_run("图 1  日常判断首页（左）与高风险优先求助反馈（右）")
    set_run_font(r, size=9, color=MUTED, italic=True)
    keep_together(caption)

    add_heading(doc, "5. 验证、可解释性与演示支撑", 1)
    add_bullet(doc, bullet_id, "规则实验室：页面内可复现 RISE、风险分级、推荐路径、个人基线和个性化排序结果。")
    add_bullet(doc, bullet_id, "验证证据板：从当前匿名档案即时计算记录数、平均恢复指数、干预覆盖率和 Safety Gate 触发数，并明确标注为演示验证。")
    add_bullet(doc, bullet_id, "自动化脚本：仓库提供 20 条规则/个性化用例、UI 冒烟测试、演示路径测试、文案安全审查、合成数据分析与提交前审查。")
    add_bullet(doc, bullet_id, "本次整理时已重新运行 20 条规则/个性化用例，结果全部通过；UI 冒烟测试脚本存在，但当前本地未安装 Playwright，未在本次会话中复跑。")

    add_heading(doc, "6. SwiftUI 原型同步情况", 1)
    ios_rows = [
        ("模块", "SwiftUI 当前实现", "说明"),
        ("匿名档案", "已实现", "多档案入口与本机隔离"),
        ("首页与判断详情", "已实现", "RISE、基线、风险、Safety Gate、推荐路径"),
        ("记录与趋势", "已实现", "多信号记录、个人基线与情绪轨迹"),
        ("陪伴与求助", "已实现", "轻干预回写、求助流程与话术"),
        ("设置与规则实验室", "已实现", "UserDefaults、证据板、15 条页面验证结果"),
        ("最终编译验证", "待完成", "当前 Windows 环境无法使用 Xcode，需在 Mac 上复核"),
    ]
    table = doc.add_table(rows=len(ios_rows), cols=3)
    fill_table_text(table, ios_rows, header=True, widths=(2500, 1900, 4960))

    add_heading(doc, "7. 当前未完成或仍属规划的能力", 1)
    add_callout(doc, "边界说明", "以下内容不应写成现有功能或已验证成果。", color=CAUTION, fill="FFF8E8")
    add_bullet(doc, bullet_id, "HealthKit 健康数据接入尚未实现，当前睡眠和步数来自手动输入或演示数据。")
    add_bullet(doc, bullet_id, "Core ML/端侧机器学习模型尚未接入，当前核心判断使用可解释规则。")
    add_bullet(doc, bullet_id, "云端同步、跨设备共享与统一身份认证未开放；当前为单机本地存储。")
    add_bullet(doc, bullet_id, "系统级后台通知未接入；设置中的提醒为原型交互状态。")
    add_bullet(doc, bullet_id, "真实用户研究结果、专业审核意见和最终演示视频仍待外部执行或录制，不能用合成数据替代。")

    doc.core_properties.title = "心晴 MindPulse 现有功能总结"
    doc.core_properties.subject = "现有功能与当前实现边界"
    doc.core_properties.author = "心晴 MindPulse 项目组"
    doc.core_properties.keywords = "MindPulse, 心晴, 功能总结, RISE, Safety Gate"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
