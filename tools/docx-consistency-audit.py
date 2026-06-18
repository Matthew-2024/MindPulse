# -*- coding: utf-8 -*-
from __future__ import annotations

import re
import sys
from dataclasses import dataclass
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"


@dataclass(frozen=True)
class Concept:
    label: str
    any_of: tuple[str, ...]


CRITICAL_DOCX: dict[str, tuple[Concept, ...]] = {
    "国赛PPT大纲.docx": (
        Concept("产品边界", ("不做诊断", "不诊断")),
        Concept("不替代咨询", ("不替代咨询", "不替代心理咨询")),
        Concept("Safety Gate", ("Safety Gate", "高风险触发后")),
        Concept("合成演示数据", ("合成演示数据",)),
        Concept("真实用户研究待回填", ("真实用户研究待", "后续用户研究计划", "如果完成用户研究")),
        Concept("规则验证", ("20 条规则", "规则/个性化用例")),
    ),
    "国赛答辩脚本.docx": (
        Concept("产品边界", ("没有做心理诊断", "不做心理诊断", "不是医学诊断")),
        Concept("不替代咨询", ("不替代专业咨询", "不替代咨询")),
        Concept("Safety Gate", ("Safety Gate", "不会继续推荐普通轻干预", "求助优先模式")),
        Concept("合成演示数据", ("合成演示数据",)),
        Concept("真实用户研究待回填", ("真实用户研究", "真实匿名试用数据")),
        Concept("专业审核待回填", ("专业审核", "待审核")),
    ),
    "国赛提交清单.docx": (
        Concept("preflight", ("npm.cmd run preflight",)),
        Concept("Safety Gate", ("Safety Gate", "高风险场景不再推荐普通轻干预")),
        Concept("合成演示数据", ("合成演示数据",)),
        Concept("真实用户研究待回填", ("没有真实试用", "补充用户研究报告", "真实结果待回填")),
        Concept("专业审核待回填", ("没有真实试用和专业审核", "补充专业老师或辅导员审核意见", "审核结论待回填")),
    ),
    "测试报告.docx": (
        Concept("preflight", ("npm.cmd run preflight",)),
        Concept("正式 docx 口径审查", ("正式 docx 口径审查", "audit:docx")),
        Concept("规则验证", ("All 20 rule case(s) passed.", "20 条规则/个性化用例")),
        Concept("合成演示数据", ("合成演示数据",)),
        Concept("A07 待执行", ("A07", "真实用户研究结果仍待")),
        Concept("A08 待执行", ("A08", "专业审核意见仍待")),
    ),
    "心理安全与伦理说明.docx": (
        Concept("产品边界", ("不是心理诊断系统", "不做诊断")),
        Concept("不替代咨询", ("不替代心理咨询",)),
        Concept("高风险求助优先", ("高风险求助优先模式", "停止普通轻干预作为主推荐路径")),
        Concept("数据最小化", ("不读取聊天记录", "不上传情绪文本", "不读取聊天记录、联系人、定位")),
    ),
    "恢复指数与风险分级算法说明.docx": (
        Concept("非医学分数", ("不是医学诊断分数", "不作为医学量表")),
        Concept("可解释", ("可解释",)),
        Concept("风险优先级", ("风险分级优先级高于恢复指数", "高风险优先级")),
        Concept("求助入口", ("求助入口",)),
    ),
    "用户研究与验证计划.docx": (
        Concept("计划口径", ("研究目标", "建议样本", "试用流程")),
        Concept("不诊断理解", ("不是心理诊断工具",)),
        Concept("匿名试用", ("匿名试用", "不收集姓名")),
    ),
    "用户研究报告模板.docx": (
        Concept("模板口径", ("完成 10-20 名同学匿名试用后", "占位内容替换为真实数据")),
        Concept("不诊断理解", ("不是心理诊断工具",)),
        Concept("隐私边界", ("不收集姓名、学号、联系方式",)),
    ),
    "专业审核访谈纪要模板.docx": (
        Concept("专业审核对象", ("心理老师", "辅导员", "校心理中心")),
        Concept("不诊断边界", ("不诊断、不替代咨询",)),
        Concept("高风险流程", ("停止普通自助建议", "打开求助入口")),
    ),
}


FORBIDDEN_PATTERNS: tuple[tuple[str, str], ...] = (
    ("诊断抑郁症", r"诊断抑郁症"),
    ("诊断焦虑症", r"诊断焦虑症"),
    ("治疗焦虑", r"治疗焦虑"),
    ("治疗抑郁", r"治疗抑郁"),
    ("治愈承诺", r"治愈"),
    ("临床验证误导", r"临床验证"),
    ("医学认证误导", r"医学认证"),
    ("已完成真实用户研究", r"已完成真实用户研究"),
    ("已完成专业审核", r"已完成专业审核"),
    ("高风险普通干预优先", r"高风险也可以先呼吸放松"),
)

ALLOWED_CONTEXTS = (
    "不要说",
    "不能说",
    "不写",
    "禁止写法",
    "不出现",
    "没有完成前",
    "不能提前写",
    "不能写",
    "不得",
    "是否可写",
    "当前不能写",
    "不代表医学认证",
    "不代表临床验证",
    "不作为诊断",
    "不是心理诊断",
    "不是医学诊断",
    "不做心理诊断",
)


def iter_document_text(doc: Document) -> list[str]:
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        parts.append(text)

    return parts


def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text)


def context_for(text: str, start: int, length: int, radius: int = 80) -> str:
    left = max(0, start - radius)
    right = min(len(text), start + length + radius)
    return normalize(text[left:right])


def is_allowed_context(context: str) -> bool:
    return any(marker in context for marker in ALLOWED_CONTEXTS)


def main() -> int:
    failures: list[str] = []
    checked_concepts = 0

    for filename, concepts in CRITICAL_DOCX.items():
        path = DOCS_DIR / filename
        if not path.exists():
            failures.append(f"Missing formal docx: {path.relative_to(ROOT)}")
            continue

        try:
            doc = Document(path)
        except Exception as exc:
            failures.append(f"Cannot read {path.relative_to(ROOT)}: {exc}")
            continue

        text = "\n".join(iter_document_text(doc))
        flat = normalize(text)

        for concept in concepts:
            checked_concepts += 1
            if not any(phrase in text for phrase in concept.any_of):
                options = " / ".join(concept.any_of)
                failures.append(f"{filename} missing {concept.label}: expected one of [{options}]")

        for label, pattern in FORBIDDEN_PATTERNS:
            for match in re.finditer(pattern, flat):
                context = context_for(flat, match.start(), len(match.group(0)))
                if not is_allowed_context(context):
                    failures.append(f"{filename} has unsafe phrase {label}: {context}")

    if failures:
        for failure in failures:
            print(f"[FAIL] {failure}", file=sys.stderr)
        print(f"\nDocx consistency audit failed with {len(failures)} issue(s).", file=sys.stderr)
        return 1

    print(
        "Docx consistency audit passed: "
        f"{len(CRITICAL_DOCX)} formal docx files and {checked_concepts} required concepts are consistent."
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
