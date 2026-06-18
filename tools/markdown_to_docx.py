from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_RE = re.compile(r"^(\s*)(\d+)\.\s+(.*)$")
UNORDERED_RE = re.compile(r"^(\s*)[-*+]\s+(.*)$")
QUOTE_RE = re.compile(r"^>\s?(.*)$")
TABLE_ALIGN_RE = re.compile(r"^\s*\|?[\s:-]+(?:\|[\s:-]+)+\|?\s*$")
INLINE_RE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")


def set_run_font(run, font_name: str, size: float | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)


def shade_element(element, fill: str) -> None:
    props = element.get_or_add_tcPr() if hasattr(element, "get_or_add_tcPr") else element.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def shade_run(run, fill: str) -> None:
    rpr = run._element.get_or_add_rPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    rpr.append(shading)


def set_cell_margins(cell, top: int, start: int, bottom: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(6)

    for level, size, color in [
        (1, 18, "27413B"),
        (2, 15, "3C6258"),
        (3, 13, "587A70"),
        (4, 12, "587A70"),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if level == 1 else 8)
        style.paragraph_format.space_after = Pt(6)


def add_inline_runs(paragraph, text: str, base_size: float = 11, color: str = "3F534D") -> None:
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, "Consolas", 10)
            run.font.color.rgb = RGBColor.from_string("3A4C47")
            shade_run(run, "EEF4F1")
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, "Microsoft YaHei", base_size)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(color)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, "Microsoft YaHei", base_size)
            run.font.color.rgb = RGBColor.from_string(color)


def split_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def consume_table(lines: list[str], index: int) -> tuple[int, list[list[str]]]:
    rows = [split_table_row(lines[index])]
    index += 2
    while index < len(lines) and "|" in lines[index]:
        rows.append(split_table_row(lines[index]))
        index += 1
    return index, rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.autofit = True

    for row_index, row in enumerate(rows):
        for col_index in range(col_count):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 90, 90, 90, 90)
            text = row[col_index] if col_index < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            add_inline_runs(p, text, 10.5, "324842")
            if row_index == 0:
                shade_element(cell._tc, "EAF5F0")
                for run in p.runs:
                    run.font.bold = True
            else:
                shade_element(cell._tc, "FCFEFD")


def add_code_block(doc: Document, code_lines: list[str]) -> None:
    for line in code_lines or [""]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.right_indent = Cm(0.4)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        shade_element(p._element, "F3F6F5")
        run = p.add_run(line)
        set_run_font(run, "Consolas", 9.5)
        run.font.color.rgb = RGBColor.from_string("35504A")


def add_list_item(doc: Document, text: str, ordered: bool, level: int) -> None:
    style_name = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.45)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(2)
    add_inline_runs(p, text)


def add_quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.right_indent = Cm(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    shade_element(p._element, "F7FAF8")
    run = p.add_run(text)
    set_run_font(run, "Microsoft YaHei", 10.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string("607A72")


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.first_line_indent = Cm(0.74)
    add_inline_runs(p, text)


def convert_markdown(md_path: Path) -> Path:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    configure_document(doc)

    title = None
    index = 0
    in_code = False
    code_lines: list[str] = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue

        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        heading = HEADING_RE.match(line)
        if heading:
            level = min(len(heading.group(1)), 4)
            content = heading.group(2).strip()
            if title is None and level == 1:
                title = content
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline_runs(p, content, 18 - (level - 1) * 2, "27413B" if level == 1 else "3C6258")
            index += 1
            continue

        if "|" in line and index + 1 < len(lines) and TABLE_ALIGN_RE.match(lines[index + 1]):
            index, rows = consume_table(lines, index)
            add_table(doc, rows)
            continue

        ordered = ORDERED_RE.match(line)
        if ordered:
            level = len(ordered.group(1).replace("\t", "    ")) // 2
            add_list_item(doc, ordered.group(3).strip(), True, level)
            index += 1
            continue

        unordered = UNORDERED_RE.match(line)
        if unordered:
            level = len(unordered.group(1).replace("\t", "    ")) // 2
            add_list_item(doc, unordered.group(2).strip(), False, level)
            index += 1
            continue

        quote = QUOTE_RE.match(line)
        if quote:
            add_quote(doc, quote.group(1).strip())
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            next_line = lines[index]
            next_stripped = next_line.strip()
            if not next_stripped:
                index += 1
                break
            if (
                HEADING_RE.match(next_line)
                or ORDERED_RE.match(next_line)
                or UNORDERED_RE.match(next_line)
                or QUOTE_RE.match(next_line)
                or next_stripped.startswith("```")
                or ("|" in next_line and index + 1 < len(lines) and TABLE_ALIGN_RE.match(lines[index + 1]))
            ):
                break
            paragraph_lines.append(next_stripped)
            index += 1
        add_paragraph(doc, " ".join(paragraph_lines))

    if in_code:
        add_code_block(doc, code_lines)

    if title is None:
        title = md_path.stem
        p = doc.add_paragraph(style="Heading 1")
        add_inline_runs(p, title, 18, "27413B")
        if not text.strip():
            empty = doc.add_paragraph()
            add_inline_runs(empty, "空白文档")

    doc.core_properties.title = title
    doc.core_properties.subject = "Markdown converted to DOCX"
    doc.core_properties.comments = "Generated from local Markdown source with preserved headings, lists, tables, quotes, and code blocks."

    out_path = md_path.with_suffix(".docx")
    doc.save(out_path)
    return out_path


def main() -> None:
    md_files = sorted(path for path in ROOT.rglob("*.md") if path.is_file())
    for md_path in md_files:
        out_path = convert_markdown(md_path)
        print(f"{md_path.relative_to(ROOT)} -> {out_path.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
